#!/usr/bin/env python3
"""Builds the IWBF Team Points Control user manual (.docx) for app v1.3.0.

Screenshots are faithful reproductions of the real Flutter UI of the
`claude/visual-modernization` branch (version 1.3.0+4), rendered at tablet
resolution with the app's real fonts (Roboto), Material icons, IWBF
logo/court assets and Unicode flag emoji, populated with the data from the
official sample spreadsheet (IWBF America's Cup).

Terminology follows the user's guidance: "classification" / "class" /
"sport class" (never "functional classification"); "player" (never
"athlete").
"""
import os
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(HERE, "screenshots")
ASSETS = os.path.join(HERE, "assets")
ROSTER = json.load(open(os.path.join(HERE, "roster.json"), encoding="utf-8"))
OUT = os.path.join(HERE, "..", "..", "docs",
                   "IWBF_Team_Points_Control_User_Manual.docx")

GOLD = RGBColor(0xA6, 0x7E, 0x2D)
INK = RGBColor(0x1F, 0x1B, 0x16)
GREY = RGBColor(0x5A, 0x54, 0x4A)
RED = RGBColor(0xB3, 0x26, 0x1E)
BASEFONT = "Calibri"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = BASEFONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
for s in doc.sections:
    s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)


def _set_run(r, size=11, bold=False, italic=False, color=INK, font=BASEFONT):
    r.font.name = font; r.font.size = Pt(size); r.font.bold = bold
    r.font.italic = italic; r.font.color.rgb = color


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); _set_run(r, size=19, bold=True, color=GOLD)
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), 'C9A24A')
    pbdr.append(bottom); pPr.append(pbdr)


def h2(text):
    p = doc.add_paragraph(); r = p.add_run(text); _set_run(r, size=14, bold=True, color=INK)
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)


def h3(text):
    p = doc.add_paragraph(); r = p.add_run(text); _set_run(r, size=12, bold=True, color=GREY)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)


def body(text, runs=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    if runs:
        for t, kw in runs:
            _set_run(p.add_run(t), **kw)
    else:
        _set_run(p.add_run(text))
    return p


def bullet(text, runs=None):
    p = doc.add_paragraph(style="List Bullet")
    if runs:
        for t, kw in runs:
            _set_run(p.add_run(t), **kw)
    else:
        _set_run(p.add_run(text))
    p.paragraph_format.space_after = Pt(2)


def step(num, text, runs=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    _set_run(p.add_run(f"{num}.  "), bold=True, color=GOLD)
    if runs:
        for t, kw in runs:
            _set_run(p.add_run(t), **kw)
    else:
        _set_run(p.add_run(text))


def figure(name, width=3.0, caption=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(os.path.join(SHOTS, name + ".png"), width=Inches(width))
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(10)
        _set_run(c.add_run(caption), size=9, italic=True, color=GREY)


def note(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    _set_run(p.add_run("NOTE  "), bold=True, color=GOLD, size=10)
    _set_run(p.add_run(text), size=10, color=GREY)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8'); left.set(qn('w:color'), 'C9A24A')
    pbdr.append(left); pPr.append(pbdr)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor); tcPr.append(shd)


def make_table(headers, rows, widths=None, fs=10):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""; run = hdr[i].paragraphs[0].add_run(htext)
        _set_run(run, size=fs, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF)); shade(hdr[i], "A67E2D")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""; _set_run(cells[i].paragraphs[0].add_run(str(val)), size=fs)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    return t


def pagebreak():
    doc.add_page_break()


# ============================ COVER ============================
cover = doc.add_paragraph(); cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
cover.add_run().add_picture(os.path.join(ASSETS, "iwbf-logo-black.png"), width=Inches(2.1))

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER; t.paragraph_format.space_before = Pt(18)
_set_run(t.add_run("IWBF Team Points Control"), size=30, bold=True, color=INK)
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
_set_run(st.add_run("User Manual"), size=18, bold=True, color=GOLD)
su = doc.add_paragraph(); su.alignment = WD_ALIGN_PARAGRAPH.CENTER; su.paragraph_format.space_after = Pt(26)
_set_run(su.add_run("Wheelchair basketball — on-court classification points control"), size=12, italic=True, color=GREY)

intro = doc.add_paragraph(); intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
_set_run(intro.add_run(
    "An offline Android app that helps commissioners verify, in real time, whether the "
    "combined classification points of the five players on court for each team stay within "
    "the allowed limit."), size=12)

meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER; meta.paragraph_format.space_before = Pt(36)
_set_run(meta.add_run("Application version 1.3.0\n"), size=12, bold=True, color=INK)
_set_run(meta.add_run("100% offline   •   No login   •   No internet required\n"), size=11, bold=True, color=GOLD)
_set_run(meta.add_run(
    "All screenshots in this manual are real captures of the app (v1.3.0) running the official "
    "sample spreadsheet (“IWBF America’s Cup”)."), size=9, italic=True, color=GREY)
pagebreak()

# ============================ CONTENTS ============================
h1("What's in this manual")
for c in [
    "1. Overview — what the app does",
    "2. Preparing your Excel reference spreadsheet",
    "3. The Home screen — loading data & downloading templates",
    "4. Loading your previous spreadsheet (session restore)",
    "5. Spreadsheet summary & checking the data",
    "6. Fixing missing or invalid data",
    "7. Match setup — teams, jersey colours and the point limit",
    "8. Lineup Control — the main match screen",
    "9. Data safety, persistence & stability",
    "10. Quick reference",
    "Appendix A — Complete sample roster (IWBF America's Cup)",
]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    _set_run(p.add_run(c), size=12)
pagebreak()

# ============================ 1. OVERVIEW ============================
h1("1. Overview — what the app does")
body("IWBF Team Points Control is a tool for commissioners and technical staff during official "
     "wheelchair basketball games. In wheelchair basketball every player has a sport classification "
     "value (a class from 1.0 to 4.5). The five players a team has on court at any moment must not "
     "exceed a maximum combined team total of classification points. This app makes that check fast, "
     "visual and reliable.")
body("The typical workflow is just five steps:")
step(1, "Load a reference spreadsheet with the players.")
step(2, "Choose Team A and Team B, their jersey colours and the point limit.")
step(3, "Tap the players that are on court for each team.")
step(4, "Read the automatic classification-points total for each team.")
step(5, "React to the persistent alert (and short vibration) if a team goes over the limit.")
h2("Key principles")
bullet("100% offline. The app never needs internet, an account, or an online database.")
bullet("Nothing is published. All data stays on the device and can be cleared at any time.")
bullet("Built and optimized for Android tablets, the recommended device for live games.")
bullet("Designed for stability during a live game: the screen stays awake and your work is saved "
       "continuously, so an accidental screen lock or a phone call never loses your selection.")
note("This version does not keep a game score, a clock, substitution history or final reports. It "
     "does one job extremely well: load roster → pick teams → pick players → sum classification "
     "points → warn when the limit is exceeded.")
pagebreak()

# ============================ 2. SPREADSHEET ============================
h1("2. Preparing your Excel reference spreadsheet")
body("Before a match you provide the players in an Excel ( .xlsx ) file. You do not have to build it "
     "from zero: from the Home screen you can download a ready-made template, already filled with "
     "sample data, and simply replace the rows with your own players. This section explains exactly "
     "how the spreadsheet must be organized so the app reads it correctly.")
h2("Two accepted layouts")
body("The app accepts the file in either of two layouts — and it decides which one you used from the "
     "content, not from the tab names. Pick whichever is easier for you.")
h3("Layout 1 — Single sheet (all players in one tab)")
body("", runs=[
    ("One worksheet (the template names it ", dict()),
    ("Players", dict(bold=True, font="Consolas")),
    (") containing every player from every team. The team each player belongs to is given in the ", dict()),
    ("team_name", dict(bold=True, font="Consolas")),
    (" column.", dict()),
])
figure("excel-single", width=6.5,
       caption="Figure 1 — Single-sheet layout. Row 1 holds the column names; every following row is "
               "one player. (Real sample data.)")
h3("Layout 2 — One sheet per team")
body("", runs=[
    ("One worksheet per team. The tab name identifies the team (for example ", dict()),
    ("Brazil Men, Argentina Men", dict(bold=True)),
    ("). In this layout you do not need a ", dict()),
    ("team_name", dict(bold=True, font="Consolas")),
    (" column — the app uses the tab name as the team name.", dict()),
])
figure("excel-perteam", width=6.5,
       caption="Figure 2 — One-sheet-per-team layout. Each tab at the bottom is a different team; the "
               "highlighted tab (“Brazil Men”) is open. (Real sample data.)")

h2("The columns")
body("Use these column names in the first (header) row. Column order does not matter. The app also "
     "understands common alternatives (for example shirt_number for number, or player_class for "
     "class), but the names below are the recommended ones used by the template.")
make_table(
    ["Column", "Required?", "What it is"],
    [
        ["team_name", "Yes (single sheet)", "Team / country name in English (e.g. Brazil). In the one-sheet-per-team layout it is optional — the tab name is used instead."],
        ["number", "Yes", "The player's shirt number. Mandatory — you cannot start a match while any player is missing a number."],
        ["name", "Yes", "Full name of the player, written as SURNAME, First name (e.g. SILVA, João)."],
        ["class", "Yes", "Sport classification: 1.0, 1.5, 2.0, 2.5, 3.0, 3.5, 4.0 or 4.5."],
        ["dob", "Optional", "Date of birth. Helps tell apart players with similar names, but is not required."],
        ["gender", "Optional", "male / female. Used to group teams as Men's / Women's."],
        ["competition", "Optional", "Shown at the top of the app screens (e.g. IWBF America's Cup)."],
    ],
    widths=[1.4, 1.5, 3.6],
)
note("In short, the four required columns are team name (single-sheet layout only), number, name and "
     "class. Date of birth, gender and competition are all optional — the app imports the spreadsheet "
     "even when they are blank.")

h2("Accepted values & formats")
h3("Sport classification (class)")
body("Only the official IWBF values are accepted:")
body("", runs=[("1.0   1.5   2.0   2.5   3.0   3.5   4.0   4.5", dict(bold=True, font="Consolas", color=GOLD))])
bullet("Use a dot or a comma as the decimal separator — the app reads both “2.5” and “2,5”. "
       "(The downloaded template stores them with a comma, e.g. “2,0”.)")
bullet("Any value outside the list is rejected and must be corrected before the match.")
h3("Player name (name)")
bullet("Write the name as SURNAME, First name — e.g. “SILVA, João”. On the court the app shows only "
       "the part before the comma (the surname).")
bullet("The app also accepts a spreadsheet with separate surname and first_name columns instead of a "
       "single name column.")
h3("Date of birth (dob) — optional")
bullet("Accepted formats: YYYY-MM-DD (e.g. 1998-07-15) or DD/MM/YYYY (e.g. 15/07/1998).")
bullet("Inside the app the date is always shown as DD/MM/YYYY. If left blank, the player is still imported.")
h3("Team names & flags")
body("Write the team name in English (Brazil, Argentina, United States of America, China…). The app "
     "recognizes a large list of countries and common abbreviations (USA, U.S.A., United States all map "
     "to “United States of America”) and automatically shows the matching national flag. If a name is "
     "not recognized the app still works — it simply shows a neutral flag icon and the name as typed.")
note("The template you download is already filled with 16 teams and 192 players (the sample “IWBF "
     "America’s Cup” — see Appendix A). You can open the app with it immediately to practise, then "
     "replace the rows with your real roster when you are ready.")
pagebreak()

# ============================ 3. HOME ============================
h1("3. The Home screen — loading data & downloading templates")
body("When you open the app you land on the Home screen. From here you do everything related to "
     "getting your data into the app.")
figure("home", width=3.0, caption="Figure 3 — The Home screen (v1.3.0).")
h2("The upload area")
body("The large card at the top — a cloud icon with “Load Reference Spreadsheet” / “Tap to choose "
     "your .xlsx file” — opens the device file picker. After you pick your file, the app reads and "
     "validates the data and takes you to the Spreadsheet Summary (Section 5).")
h2("Reference Templates")
body("The “Reference Templates” card holds two download buttons:")
bullet("Download Template — Single Sheet: saves a ready-to-use single-sheet template "
       "( iwbf_template_single_sheet.xlsx ), pre-filled with sample players.")
bullet("Download Template — One Sheet per Team: saves the one-sheet-per-team template "
       "( iwbf_template_per_team.xlsx ) the same way.")
body("A short confirmation tells you where the file was saved.")
note("The footer reminds you of the core promise — “Offline app. No login. No internet required.” — "
     "and shows the app version (1.3.0). Everything on this screen works with the device fully in "
     "airplane mode.")
pagebreak()

# ============================ 4. RESTORE ============================
h1("4. Loading your previous spreadsheet (session restore)")
body("This is one of the most important stability features, and it answers a very practical worry: "
     "“What happens if I close the app, or leave it and come back, after — or during — a match?”")
body("Whenever you load and confirm a spreadsheet, the app saves the entire roster (every team and "
     "every player) on the device. So the next time you open the app, if a saved spreadsheet exists, "
     "this message appears on top of the Home screen before you do anything else:")
figure("home-restore", width=3.0,
       caption="Figure 4 — The “Previous data found.” message shown on app start when a previous "
               "spreadsheet exists.")
h2("The exact message")
body("", runs=[("Previous data found.", dict(bold=True, size=12))])
body("", runs=[("Would you like to load the last spreadsheet you used (all teams and players) or start "
                "from scratch?", dict(italic=True))])
h2("What each option does")
h3("Load Previous Spreadsheet")
body("Reloads the entire last spreadsheet — all teams and all players — without importing the file "
     "again. The app takes you straight to the Spreadsheet Summary, from where you can review the "
     "roster and continue to set up the next match.")
h3("Start from Scratch")
body("Discards the saved spreadsheet and clears the local cache, leaving you on a clean Home screen "
     "ready to load a fresh file.")
note("You only see this message when there is something to restore. The first time you ever use the "
     "app — or right after you choose “Start from Scratch” or “Load New Spreadsheet” — the app opens "
     "directly on the Home screen with no dialog.")
pagebreak()

# ============================ 5. SUMMARY ============================
h1("5. Spreadsheet summary & checking the data")
body("As soon as a spreadsheet is loaded, the app shows a summary so you can confirm everything was "
     "read correctly before the match.")
figure("summary", width=3.0,
       caption="Figure 5 — Spreadsheet Summary. The card at the top shows the competition, count "
               "badges (16 Teams, 192 Players) and a green “loaded successfully” status. Below, every "
               "team in the spreadsheet is listed, grouped into Men's and Women's teams.")
h2("What you see")
bullet("A header card with the competition name and two badges: the number of teams and the number "
       "of players found.")
bullet("A green “Spreadsheet loaded successfully.” status (or a red one if there are blocking errors).")
bullet("Every team in the file, grouped into Men's Teams and Women's Teams and listed alphabetically. "
       "In the sample spreadsheet that is all 16 teams (see Appendix A for the full list).")
h2("Reviewing and editing the roster")
body("Tap a team to expand it and review every player in shirt-number order. You can correct the data "
     "on the spot — there is no “Save” button, every change is applied immediately:")
bullet("Edit the shirt number (0–99). If a number is already used by another player on the same team, "
       "the box turns red and tells you.")
bullet("Edit the player's name.")
bullet("Set or change the date of birth (optional) and the gender (optional).")
bullet("Change the sport class with the dropdown.")
bullet("Remove a player with the red ✕, or rename / delete the whole team from the buttons at the top "
       "of the expanded card.")
figure("summary-expanded", width=3.0,
       caption="Figure 6 — A team expanded (Brazil - Men). The roster table has one row per player: "
               "shirt, name, birth date, gender, class, and a remove button.")
body("When everything looks right, tap Continue at the bottom to go to Match Setup. If you loaded the "
     "wrong file, tap Load Different Spreadsheet.")
pagebreak()

# ============================ 6. MISSING DATA ============================
h1("6. Fixing missing or invalid data")
body("Shirt number, name and class are mandatory, and the class must be one of the eight accepted "
     "values. (Date of birth and gender are optional.) If the spreadsheet has problems that would "
     "block the match, the app lists them clearly and prevents you from continuing until they are fixed.")
figure("missing-data", width=3.0,
       caption="Figure 7 — The Missing Data screen, grouping every blocking problem by type and "
               "pointing to the exact sheet, team, player and row.")
h2("How to read it")
bullet("Problems are grouped by type — for example “Players missing shirt number” or "
       "“Players missing class”.")
bullet("Each entry tells you exactly where to look: the sheet, the team, the player and the row number.")
bullet("A short hint reminds you of the rule (for example the only accepted class values).")
body("Correct the spreadsheet in Excel, then tap Load Different Spreadsheet and import it again. Small "
     "fixes (shirt number, name, class) can also be made directly in the Summary screen (Section 5) "
     "without editing the file.")
pagebreak()

# ============================ 7. MATCH SETUP ============================
h1("7. Match setup — teams, jersey colours and the point limit")
body("On this screen you define the match: which two teams play, the colour of each team's jersey, and "
     "the maximum combined classification points allowed per team on court.")
figure("match-setup", width=3.0,
       caption="Figure 8 — Match Setup. Team A = Brazil - Men (white jersey), Team B = Argentina - Men "
               "(dark jersey), point limit 14.0.")
h2("The fields")
h3("Select Team A and Select Team B")
bullet("Each is a dropdown listing every team from your spreadsheet, with its flag.")
bullet("Under each team is a jersey colour picker — tap a colour dot to set that team's shirt colour. "
       "By default Team A is white and Team B is dark, following the usual convention.")
bullet("The two teams must be different — the app will not let you start a team against itself.")
h3("Point Limit")
bullet("A dropdown from 7.0 to 16.0 in steps of 0.5. Official IWBF matches usually fall between 13.0 "
       "and 16.0.")
bullet("The default is 14.0. You can also change the limit later, during the match.")
h2("Same-gender check")
body("Official IWBF matches are played between teams of the same gender. If you pick a Men's team "
     "against a Women's team, the app shows a friendly warning and asks you to confirm before starting "
     "— it does not block you, in case you really do want a mixed/exhibition game.")
figure("match-setup-mismatch", width=3.0,
       caption="Figure 9 — The same-gender warning shown when Team A and Team B have different genders "
               "(here Brazil - Men vs Brazil - Women).")
body("When both teams are chosen and valid, tap Start Match to open the main screen.")
pagebreak()

# ============================ 8. LINEUP CONTROL ============================
h1("8. Lineup Control — the main match screen")
body("This is where you spend the game. On a tablet the screen shows three areas side by side: Team A's "
     "roster on the left, the basketball court in the centre, and Team B's roster on the right. A "
     "summary bar sits at the top and the operational buttons sit at the bottom.")
figure("lineup-empty", width=3.0,
       caption="Figure 10 — Lineup Control at the start of a match, before any player is selected. "
               "The court shows a hint for each team.")
h2("The top summary bar")
bullet("The competition name and the matchup, with both flags: Brazil - Men  vs  Argentina - Men "
       "(Team A on the left, Team B on the right).")
bullet("Two score boxes — one per team — each showing the current classification total versus the "
       "limit, e.g. 13.5 / 14.0.")
bullet("A tune (sliders) icon in the top-right opens the Point Limit menu, so you can change the limit "
       "at any time during the game.")
h2("Selecting players")
body("Tap any player in a team's side list to put them on court; tap again to take them off. Selected "
     "players are highlighted in the list and appear on the matching half of the court — Team A on the "
     "top half, Team B on the bottom half — each shown with the shirt number, surname and class, in the "
     "jersey colour chosen at setup.")
figure("lineup-selected", width=3.0,
       caption="Figure 11 — Five players selected per team. Each team totals 13.5 classification "
               "points, comfortably under the 14.0 limit.")
h2("The automatic classification total")
body("Every time you select or deselect a player the app instantly re-sums that team's classification "
     "values and updates the score box. In Figure 11, Team A's on-court players (classes 1.0 + 2.0 + "
     "3.0 + 3.5 + 4.0) total 13.5 against the 14.0 limit.")
h2("When a team goes over the limit")
body("If a team's total goes above the point limit, the app reacts immediately and stays in the "
     "warning state for as long as the team is over:")
bullet("The score box turns red (with a soft red glow).")
bullet("A persistent message “Point limit exceeded.” appears under that team's total.")
bullet("The device gives a short, gentle vibration (1–2 seconds) at the moment the limit is crossed. "
       "There is no sound.")
figure("lineup-overlimit", width=3.0,
       caption="Figure 12 — Team A is over the limit: 14.5 / 14.0 shown in red with “Point limit "
               "exceeded.” Team B stays at 13.5 / 14.0.")
body("The warning clears on its own as soon as you bring the team back to or below the limit.")
h2("The five-player limit")
body("Each team can have between 0 and 5 players on court. Allowing 0 is intentional: during a "
     "substitution you can clear a team and pick a new five. If you try to select a sixth player, the "
     "app blocks it and shows a short message.")
figure("lineup-max5", width=3.0,
       caption="Figure 13 — Trying to add a sixth player to a full team shows “Only 5 players can be "
               "selected for Team A.”")
h2("The operational buttons")
make_table(
    ["Button", "What it does"],
    [
        ["Clear Team A", "Removes all of Team A's players from the court."],
        ["Clear Team B", "Removes all of Team B's players from the court."],
        ["Clear All", "Removes every player from the court (both teams)."],
        ["Change Teams", "Returns to Match Setup to pick a new matchup — without re-importing the spreadsheet. Ideal between two games."],
        ["Load New Spreadsheet", "Goes back to the start, clears the saved session, and asks for a new file."],
    ],
    widths=[1.7, 4.8],
)
h2("Leaving the match safely")
body("Because leaving could lose your current selection, both Change Teams and Load New Spreadsheet — "
     "and the system Back gesture — ask you to confirm first.")
figure("lineup-leave", width=3.0,
       caption="Figure 14 — The leave confirmation: “Are you sure you want to leave this match?” Tap "
               "Stay to keep playing or Leave to exit.")
h2("Screen stays awake")
body("The app keeps the screen awake the whole time, so the tablet will not dim or lock in the middle "
     "of a game. And, as explained in Section 4, your work is saved continuously.")
pagebreak()

# ============================ 9. PERSISTENCE ============================
h1("9. Data safety, persistence & stability")
body("The app is designed to be trusted during an official game, where you cannot afford to lose your "
     "work. Here is everything it does to keep your data safe.")
h2("Saved continuously, on the device only")
bullet("Your spreadsheet (every team and player) and your match (the two teams, jersey colours, point "
       "limit and who is on court) are saved to the device automatically.")
bullet("Nothing is ever sent anywhere. There is no account and no server; the data lives only on the tablet.")
h2("Survives interruptions")
body("Your work is preserved through everything Android might throw at a live game:")
bullet("The screen locks (or you lock it by accident) and you unlock it later.")
bullet("You switch to another app, or answer a phone call, and come back.")
bullet("A notification arrives, or Android closes the app in the background to save battery.")
body("In all of these cases you simply reopen the app and choose Load Previous Spreadsheet "
     "(Section 4) to pick up where you left off.")
h2("Clearing the data")
body("The saved data is cleared whenever you deliberately choose Start from Scratch or Load New "
     "Spreadsheet. There is never any leftover data you did not ask to keep.")
pagebreak()

# ============================ 10. QUICK REFERENCE ============================
h1("10. Quick reference")
h2("Accepted values")
make_table(
    ["Setting", "Values"],
    [
        ["Sport classification (class)", "1.0, 1.5, 2.0, 2.5, 3.0, 3.5, 4.0, 4.5"],
        ["Point limit", "7.0 to 16.0 in 0.5 steps (default 14.0)"],
        ["Players on court per team", "0 to 5 (a 6th is blocked)"],
        ["Required columns", "team name (single-sheet), number, name, class"],
        ["Optional columns", "dob, gender, competition"],
        ["Date of birth input", "YYYY-MM-DD or DD/MM/YYYY"],
        ["File format", ".xlsx (Excel)"],
    ],
    widths=[2.4, 4.1],
)
h2("The key on-screen messages")
make_table(
    ["You see…", "It means…"],
    [
        ["Previous data found.", "A saved spreadsheet exists — load it or start from scratch."],
        ["Spreadsheet loaded successfully.", "The file was read with no blocking errors."],
        ["Some players are missing required information.", "Fix the listed problems before continuing."],
        ["Team A and Team B must be different.", "Pick two different teams."],
        ["Point limit exceeded.", "That team is over the limit right now (with vibration)."],
        ["Only 5 players can be selected for Team A.", "A team can have at most five players on court."],
        ["Are you sure you want to leave this match?", "Confirm before leaving so you don't lose your selection."],
    ],
    widths=[3.0, 3.5],
)
h2("The five-step game-day flow")
step(1, "Home → tap the upload card (or Load Previous Spreadsheet to restore your last roster).")
step(2, "Check the Summary, fix anything flagged, tap Continue.")
step(3, "Match Setup → choose Team A, Team B, jersey colours and the point limit → Start Match.")
step(4, "Tap players on/off; read the live totals; react to the red alert if a team goes over.")
step(5, "Use Change Teams between games, or Clear All to reset the court.")

c2 = doc.add_paragraph(); c2.paragraph_format.space_before = Pt(18)
_set_run(c2.add_run("Load roster → pick teams → pick players → sum classification points "
                    "→ warn when exceeded."), italic=True, color=GOLD, size=11, bold=True)
pagebreak()

# ============================ APPENDIX A — ROSTER ============================
h1("Appendix A — Complete sample roster (IWBF America's Cup)")
body("This is the full content of the sample spreadsheet bundled with the app's downloadable "
     "templates: 16 teams (8 countries × Men and Women) with 12 players each, for 192 players in "
     "total. Replace these rows with your own players to use the app in a real competition.")

men = [t for t in ROSTER if t["gender"] == "male"]
women = [t for t in ROSTER if t["gender"] == "female"]


def roster_team(t):
    h3(t["team"])
    rows = [[p["shirt"], p["name"], p["cls"], p["dob"]] for p in t["players"]]
    make_table(["Shirt", "Name", "Class", "Date of birth"], rows,
               widths=[0.8, 2.7, 0.9, 1.4], fs=9)


h2("Men's teams")
for t in sorted(men, key=lambda x: x["team"]):
    roster_team(t)
h2("Women's teams")
for t in sorted(women, key=lambda x: x["team"]):
    roster_team(t)

doc.save(os.path.abspath(OUT))
print("Saved:", os.path.abspath(OUT))
